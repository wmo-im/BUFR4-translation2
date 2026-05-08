#!/usr/bin/env python3
"""
Generate Word (.docx) tables from BUFR CSV files.

Uses table_config.json for language-specific titles, headers, and styling.

Usage:
    python scripts/csv2docx.py <lang> [--table B|D|CF|all] [--class NN] [--outdir DIR]

Examples:
    python scripts/csv2docx.py ru                          # all tables, all classes
    python scripts/csv2docx.py ru --table B                # all Table B classes
    python scripts/csv2docx.py ru --table D --class 07     # Table D class 07 only
    python scripts/csv2docx.py fr --table CF --class 20    # French CodeFlag class 20
"""

import argparse, csv, json, os, re, sys, glob

from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# ── Config loader ────────────────────────────────────────────────────

def load_config():
    """Load table_config.json from the scripts directory."""
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'table_config.json')
    with open(config_path, encoding='utf-8') as f:
        return json.load(f)


def get_lang_config(config, lang, table_type):
    """Get language-specific config for a table type, falling back to English."""
    langs = config['languages']
    lc = langs.get(lang, langs.get('en', {}))
    return lc.get(table_type, {})


# ── Notes loader ─────────────────────────────────────────────────────

def load_notes(lang, base_dir):
    """Load notes for a language. No English fallback."""
    notes_dir = os.path.join(base_dir, 'notes')
    all_notes = {}
    for prefix in ['BUFRCREX_CodeFlag', 'BUFRCREX_TableB', 'BUFR_TableC', 'BUFR_TableD']:
        nf = os.path.join(notes_dir, f'{prefix}_notes_{lang}.csv')
        if os.path.exists(nf):
            with open(nf, encoding='utf-8-sig') as f:
                for r in csv.DictReader(f):
                    nid = r.get('noteID', '').strip()
                    text = r.get(f'note_{lang}', r.get('note', '')).strip()
                    if nid and text:
                        all_notes[nid] = text
    return all_notes


def resolve_note(row, lang, notes_db):
    """Resolve noteIDs to full note text."""
    note_text = row.get(f'Note_{lang}', '').strip()
    note_ids = row.get('noteIDs', '').strip()
    if not note_ids:
        return note_text
    resolved = []
    for nid in note_ids.split(','):
        nid = nid.strip()
        if nid in notes_db:
            resolved.append(f'Note {nid}: {notes_db[nid]}')
    return ' | '.join(resolved) if resolved else note_text


# ── FXY formatting ───────────────────────────────────────────────────

def format_fxy(code, fmt='compact'):
    """Format FXY. compact: 002001 as-is. spaced: 002001 → 0 02 001."""
    c = code.strip()
    if fmt == 'spaced' and len(c) == 6 and c.isdigit():
        return f'{c[0]} {c[1:3]} {c[3:]}'
    return c


# ── Unit superscript formatting ──────────────────────────────────────

# Pattern: letter(s) followed by digit(s) that are exponents, e.g. m2, s-2, m-1
# Also handles patterns like: kg m-2 s-1, m2 s-2, W m-2 sr-1
# Match lowercase letter + exponent (m2, s-1, sr-1, deg2) but NOT uppercase acronyms (IA5)
UNIT_EXP_RE = re.compile(r'(?<![A-Z])([a-z°])(-?\d+)(?=\s|$|[·⋅/)])')


def set_cell_with_units(cell, text, style, is_unit_col=False):
    """Set cell text, with superscript formatting for unit exponents if is_unit_col."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

    clean = ''.join(c for c in str(text or '') if ord(c) >= 32 or c in '\t\n\r')

    if not is_unit_col or not UNIT_EXP_RE.search(clean):
        run = p.add_run(clean)
        run.font.name = style['font']
        run.font.size = style['data_size']
        return

    # Split text into segments: normal text + superscript exponents
    pos = 0
    for m in UNIT_EXP_RE.finditer(clean):
        # Text before the match (including the base letter)
        before = clean[pos:m.start() + 1]  # include the letter
        exp = m.group(2)  # the exponent digits

        if before:
            run = p.add_run(before)
            run.font.name = style['font']
            run.font.size = style['data_size']

        # Superscript exponent
        run = p.add_run(exp)
        run.font.name = style['font']
        run.font.size = style['data_size']
        run.font.superscript = True

        pos = m.end()

    # Remaining text after last match
    if pos < len(clean):
        run = p.add_run(clean[pos:])
        run.font.name = style['font']
        run.font.size = style['data_size']


# ── Global notes loader ─────────────────────────────────────────────

def load_global_notes(lang, base_dir, table_type, class_or_category=None):
    """Load global notes for a specific class/category.

    For Table B/D: class_or_category = class number like "02"
    For CodeFlag: notes are per-FXY, loaded separately via load_codeflag_fxy_notes()
    """
    notes_dir = os.path.join(base_dir, 'notes')

    prefix_map = {
        'table_b': 'BUFRCREX_TableB',
        'table_d': 'BUFR_TableD',
        'codeflag': 'BUFRCREX_CodeFlag',
    }
    prefix = prefix_map.get(table_type, '')
    if not prefix:
        return []

    # Read global noteIDs — filter by tableNo matching our class/category
    global_file = os.path.join(notes_dir, f'{prefix}_global_noteIDs.csv')
    if not os.path.exists(global_file):
        return []

    relevant_ids = set()
    with open(global_file, encoding='utf-8-sig') as f:
        for r in csv.DictReader(f):
            tno = r.get('tableNo', '').strip()
            nid = r.get('noteID', '').strip()
            if not nid:
                continue
            if class_or_category is None:
                relevant_ids.add(nid)
            elif tno == class_or_category:
                relevant_ids.add(nid)

    if not relevant_ids:
        return []

    # Read notes file for this language
    notes_file = os.path.join(notes_dir, f'{prefix}_notes_{lang}.csv')
    if not os.path.exists(notes_file):
        return []

    notes_text = []
    with open(notes_file, encoding='utf-8-sig') as f:
        for r in csv.DictReader(f):
            nid = r.get('noteID', '').strip()
            text = r.get(f'note_{lang}', r.get('note', '')).strip()
            if nid in relevant_ids and text:
                notes_text.append(f'Note {nid}: {text}')

    return notes_text


def load_codeflag_fxy_notes(lang, base_dir, fxy):
    """Load notes for a specific CodeFlag FXY (e.g., '0 02 003' or '002003')."""
    # Normalize FXY to spaced format for matching
    fxy_spaced = fxy.strip()
    if len(fxy_spaced) == 6 and fxy_spaced.isdigit():
        fxy_spaced = f'{fxy_spaced[0]} {fxy_spaced[1:3]} {fxy_spaced[3:]}'

    return load_global_notes(lang, base_dir, 'codeflag', class_or_category=fxy_spaced)


# ── Styling helpers ──────────────────────────────────────────────────

def apply_style(config):
    """Extract style settings from config."""
    s = config['style']
    return {
        'font': s['font'],
        'title_size': Pt(s['title_size']),
        'header_size': Pt(s['header_size']),
        'data_size': Pt(s['data_size']),
        'header_bg': s['header_bg'],
        'header_bold': s['header_bold'],
        'border_color': s['border_color'],
        'border_size': s['border_size'],
        'margin_top': s['cell_margin_top'],
        'margin_bottom': s['cell_margin_bottom'],
        'margin_left': s['cell_margin_left'],
        'margin_right': s['cell_margin_right'],
    }


def set_cell(cell, text, style, bold=False, is_header=False):
    """Set cell text with config-driven styling."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    # Sanitize control chars
    clean = ''.join(c for c in str(text or '') if ord(c) >= 32 or c in '\t\n\r')
    run = p.add_run(clean)
    run.font.name = style['font']
    run.font.size = style['header_size'] if is_header else style['data_size']
    run.font.bold = bold


def style_table(table, style):
    """Apply border and cell margin styling to the entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})

    # Borders
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        borders.append(borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): str(style['border_size']),
            qn('w:space'): '0',
            qn('w:color'): style['border_color'],
        }))
    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    # Cell margins
    margins = tblPr.makeelement(qn('w:tblCellMar'), {})
    for side, val in [('top', style['margin_top']), ('left', style['margin_left']),
                      ('bottom', style['margin_bottom']), ('right', style['margin_right'])]:
        margins.append(margins.makeelement(qn(f'w:{side}'), {
            qn('w:w'): str(val), qn('w:type'): 'dxa'}))
    existing = tblPr.find(qn('w:tblCellMar'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(margins)


def add_header_row(table, headers, style):
    """Style the header row."""
    row = table.rows[0]
    # Mark as repeating header
    trPr = row._tr.get_or_add_trPr()
    trPr.append(trPr.makeelement(qn('w:tblHeader'), {}))

    for cell, hdr in zip(row.cells, headers):
        set_cell(cell, hdr, style, bold=style['header_bold'], is_header=True)
        # Shading
        tcPr = cell._element.get_or_add_tcPr()
        tcPr.append(tcPr.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto',
            qn('w:fill'): style['header_bg']}))


def set_cant_split(row):
    """Prevent row from breaking across pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(trPr.makeelement(qn('w:cantSplit'), {}))


def set_landscape(doc):
    """Set page to landscape A4."""
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width = Cm(29.7)
    s.page_height = Cm(21.0)
    s.left_margin = Cm(1.5)
    s.right_margin = Cm(1.5)
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)


def add_title(doc, text, style):
    """Add centered bold title."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = style['font']
    run.font.size = style['title_size']
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Generic row builder ──────────────────────────────────────────────

def get_row_values(row, column_map, lang, notes_db, fxy_fmt='compact'):
    """Extract values from a CSV row using the column map."""
    fxy_cols = {'FXY', 'FXY1', 'FXY2'}
    vals = []
    for col_template in column_map:
        col = col_template.replace('{lang}', lang)
        val = row.get(col, '')
        if col_template in fxy_cols:
            val = format_fxy(val, fxy_fmt)
        if col_template == 'Note_{lang}':
            val = resolve_note(row, lang, notes_db)
        vals.append(val)
    return vals


# ── Table B Generator ────────────────────────────────────────────────

def generate_table_b(lang, csv_path, doc, notes_db, config):
    style = apply_style(config)
    lc = get_lang_config(config, lang, 'table_b')
    headers = lc.get('headers', [])
    col_map = lc.get('column_map', [])

    with open(csv_path, encoding='utf-8-sig') as f:
        rows = list(csv.DictReader(f))
    if not rows:
        return

    class_no = rows[0].get('ClassNo', '')
    class_name = rows[0].get(f'ClassName_{lang}', '')
    title = lc.get('title', 'Class {class_no}: {class_name}').format(
        class_no=class_no, class_name=class_name)
    add_title(doc, title, style)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    style_table(table, style)
    add_header_row(table, headers, style)

    # Identify unit columns for superscript formatting
    unit_cols = {i for i, c in enumerate(col_map)
                 if 'Unit' in c or 'UNIT' in c.upper()}

    fxy_fmt = lc.get('fxy_format', 'spaced')
    for r in rows:
        row_cells = table.add_row().cells
        set_cant_split(table.rows[-1])
        vals = get_row_values(r, col_map, lang, notes_db, fxy_fmt)
        for i, (cell, val) in enumerate(zip(row_cells, vals)):
            if i in unit_cols:
                set_cell_with_units(cell, val, style, is_unit_col=True)
            else:
                set_cell(cell, val, style)

    # Global notes at the end — only for THIS class
    base_dir = os.path.dirname(os.path.dirname(csv_path))
    global_notes = load_global_notes(lang, base_dir, 'table_b', class_or_category=class_no)
    if global_notes:
        doc.add_paragraph()
        for note in global_notes:
            p = doc.add_paragraph()
            run = p.add_run(note)
            run.font.name = style['font']
            run.font.size = Pt(7)
            run.font.italic = True


# ── Table D Generator ────────────────────────────────────────────────

def generate_table_d(lang, csv_path, doc, notes_db, config):
    style = apply_style(config)
    lc = get_lang_config(config, lang, 'table_d')
    # Sub-table headers: Table references, Element name, Element description, Note
    sub_headers = lc.get('sub_headers', ['TABLE REFERENCES', 'ELEMENT NAME',
                                          'ELEMENT DESCRIPTION', 'NOTE'])
    sub_col_map = lc.get('sub_column_map', ['FXY2', 'ElementName_{lang}',
                                             'ElementDescription_{lang}', 'Note_{lang}'])

    with open(csv_path, encoding='utf-8-sig') as f:
        rows = list(csv.DictReader(f))
    if not rows:
        return

    cat = rows[0].get('Category', '')
    cat_name = rows[0].get(f'CategoryOfSequences_{lang}', '')
    title = lc.get('title', 'Category {category}: {category_name}').format(
        category=cat, category_name=cat_name)
    add_title(doc, title, style)

    # Group rows by FXY1
    fxy1_groups = []
    current_fxy1 = None
    current_group = []
    for r in rows:
        fxy1 = r.get('FXY1', '')
        if fxy1 != current_fxy1:
            if current_group:
                fxy1_groups.append((current_fxy1, current_group))
            current_fxy1 = fxy1
            current_group = [r]
        else:
            current_group.append(r)
    if current_group:
        fxy1_groups.append((current_fxy1, current_group))

    fxy_fmt = lc.get('fxy_format', 'spaced')

    # Create one sub-table per FXY1 sequence
    for fxy1, group in fxy1_groups:
        title_text = group[0].get(f'Title_{lang}', '')
        subtitle_text = group[0].get(f'SubTitle_{lang}', '')

        # FXY1 heading: "3 02 001 (Title)" or "3 02 001 Title" if already parenthesized
        heading = format_fxy(fxy1, fxy_fmt)
        if title_text:
            if title_text.startswith('('):
                heading += f' {title_text}'
            else:
                heading += f' ({title_text})'
        p = doc.add_paragraph()
        run = p.add_run(heading)
        run.font.name = style['font']
        run.font.size = Pt(9)
        run.font.bold = True

        # Sub-table
        table = doc.add_table(rows=1, cols=len(sub_headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        style_table(table, style)
        add_header_row(table, sub_headers, style)

        for r in group:
            vals = get_row_values(r, sub_col_map, lang, notes_db, fxy_fmt)
            row_cells = table.add_row().cells
            set_cant_split(table.rows[-1])
            for cell, val in zip(row_cells, vals):
                set_cell(cell, val, style)

        doc.add_paragraph()  # spacing between sequences

    # Global notes at the end — only for THIS category
    base_dir = os.path.dirname(os.path.dirname(csv_path))
    global_notes = load_global_notes(lang, base_dir, 'table_d', class_or_category=cat)
    if global_notes:
        doc.add_paragraph()
        for note in global_notes:
            p = doc.add_paragraph()
            run = p.add_run(note)
            run.font.name = style['font']
            run.font.size = Pt(7)
            run.font.italic = True


# ── CodeFlag Generator ───────────────────────────────────────────────

def generate_codeflag(lang, csv_path, doc, notes_db, config):
    style = apply_style(config)
    lc = get_lang_config(config, lang, 'codeflag')

    with open(csv_path, encoding='utf-8-sig') as f:
        rows = list(csv.DictReader(f))
    if not rows:
        return

    # Group by FXY
    fxy_groups = {}
    for r in rows:
        fxy = r.get('FXY', '')
        fxy_groups.setdefault(fxy, []).append(r)

    cls = os.path.basename(csv_path).split('_')[-1].replace('.csv', '')
    title = lc.get('title', 'Class {class_no}').format(class_no=cls)
    add_title(doc, title, style)

    for fxy, group in fxy_groups.items():
        elem_name = group[0].get(f'ElementName_{lang}', '')

        # Sub-heading
        fxy_fmt = lc.get('fxy_format', 'compact')
        p = doc.add_paragraph()
        run = p.add_run(f'{format_fxy(fxy, fxy_fmt)}  {elem_name}')
        run.font.name = style['font']
        run.font.size = Pt(9)
        run.font.bold = True

        # Determine column variant
        has_sub1 = any(r.get(f'EntryName_sub1_{lang}', '').strip() for r in group)
        has_sub2 = any(r.get(f'EntryName_sub2_{lang}', '').strip() for r in group)

        if has_sub2:
            headers = lc.get('headers_sub2', lc['headers'])
            col_map = lc.get('column_map_sub2', lc['column_map'])
        elif has_sub1:
            headers = lc.get('headers_sub1', lc['headers'])
            col_map = lc.get('column_map_sub1', lc['column_map'])
        else:
            headers = lc['headers']
            col_map = lc['column_map']

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        style_table(table, style)

        # CodeFlag sub-tables: no tblHeader repeat (tables are small)
        # Use add_header_row but WITHOUT tblHeader flag
        row = table.rows[0]
        for cell, hdr in zip(row.cells, headers):
            set_cell(cell, hdr, style, bold=style['header_bold'], is_header=True)
            tcPr = cell._element.get_or_add_tcPr()
            tcPr.append(tcPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear', qn('w:color'): 'auto',
                qn('w:fill'): style['header_bg']}))

        # Narrow Code Figure column (first column)
        for row_obj in table.rows:
            row_obj.cells[0].width = Cm(2.0)

        for r in group:
            vals = get_row_values(r, col_map, lang, notes_db, fxy_fmt)
            row_cells = table.add_row().cells
            set_cant_split(table.rows[-1])
            for cell, val in zip(row_cells, vals):
                set_cell(cell, val, style)

        # Notes for THIS FXY — after its sub-table
        base_dir = os.path.dirname(os.path.dirname(csv_path))
        fxy_notes = load_codeflag_fxy_notes(lang, base_dir, fxy)
        if fxy_notes:
            for note in fxy_notes:
                p = doc.add_paragraph()
                run = p.add_run(note)
                run.font.name = style['font']
                run.font.size = Pt(7)
                run.font.italic = True

        doc.add_paragraph()  # spacing between FXY tables


# ── Main ─────────────────────────────────────────────────────────────

def find_files(lang, table_type, class_no, base_dir):
    lang_dir = os.path.join(base_dir, {
        'en': 'english', 'fr': 'french', 'es': 'spanish', 'ru': 'russian'
    }.get(lang, lang))

    patterns = {
        'B': f'BUFRCREX_TableB_{lang}_*.csv',
        'D': f'BUFR_TableD_{lang}_*.csv',
        'CF': f'BUFRCREX_CodeFlag_{lang}_*.csv',
    }

    types = ['B', 'D', 'CF'] if table_type == 'all' else [table_type]
    files = []
    for t in types:
        matched = sorted(glob.glob(os.path.join(lang_dir, patterns[t])))
        if class_no:
            matched = [f for f in matched if f'_{class_no}.' in f or f'_{class_no.zfill(2)}.' in f]
        files.extend([(t, f) for f in matched])
    return files


def main():
    parser = argparse.ArgumentParser(description='Generate Word tables from BUFR CSVs')
    parser.add_argument('lang', help='Language code (ru, fr, es, en)')
    parser.add_argument('--table', default='all', choices=['B', 'D', 'CF', 'all'])
    parser.add_argument('--class', dest='class_no', default=None)
    parser.add_argument('--outdir', default=None)
    args = parser.parse_args()

    config = load_config()
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    out_dir = args.outdir or os.path.join(base_dir, 'docx', args.lang)
    os.makedirs(out_dir, exist_ok=True)

    files = find_files(args.lang, args.table, args.class_no, base_dir)
    if not files:
        print(f'No files found for lang={args.lang} table={args.table} class={args.class_no}')
        sys.exit(1)

    notes_db = load_notes(args.lang, base_dir)
    print(f'Loaded {len(notes_db)} notes for {args.lang}')

    generators = {
        'B': generate_table_b,
        'D': generate_table_d,
        'CF': generate_codeflag,
    }

    total = 0
    for table_type, csv_path in files:
        doc = Document()
        set_landscape(doc)
        basename = os.path.splitext(os.path.basename(csv_path))[0]
        generators[table_type](args.lang, csv_path, doc, notes_db, config)
        doc.save(os.path.join(out_dir, f'{basename}.docx'))
        total += 1
        print(f'  {basename}.docx')

    print(f'\n{total} files written to {out_dir}')


if __name__ == '__main__':
    main()
